import os
import re
import json
from docx import Document

def get_docx_text(path):
    """Extracts text from paragraphs AND tables to ensure nothing is missed."""
    doc = Document(path)
    content = []
    # Extract from paragraphs
    for para in doc.paragraphs:
        content.append(para.text)
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Add newline to keep structure
                content.append(cell.text + "\n")
    return "\n".join(content)

def clean_text(text):
    """Removes tags and standardizes whitespace."""
    if not text: return ""
    text = re.sub(r'\\', '', text)
    # Convert multiple spaces/newlines into clean structure
    return text.strip()

def parse_nrc_batch(raw_text):
    """Parses raw text using multiple possible separators."""
    # Split by common separators (case-insensitive)
    # Handles "End of Document", "END OF DOCUMENT", or "Document 1 of 250"
    articles = re.split(r'(?i)End of Document|Document \d+ of \d+', raw_text)
    
    extracted = []
    for art in articles:
        if "NRC" not in art or "VOLLEDIGE TEKST" not in art:
            continue
            
        # 1. Title: Look for text followed by NRC
        # We look for the title at the start of the block or after empty lines
        title_match = re.search(r'(?:^|\n)\s*(\d+\.\s*)?([^\n]+)\n\s*NRC', art, re.DOTALL)
        title = clean_text(title_match.group(2)) if title_match else "Unknown Title"

        # 2. Date: The line immediately following NRC
        date_match = re.search(r'NRC\s*\n\s*([^\n]+)', art)
        date = clean_text(date_match.group(1)) if date_match else "Unknown Date"

        # 3. Author: Byline or footer
        author_match = re.search(r'Byline:\s*([^\n]+)', art)
        author = clean_text(author_match.group(1)) if author_match else "Redactie"

        # 4. Body: Everything between VOLLEDIGE TEKST and footers
        body_match = re.search(r'VOLLEDIGE TEKST:(.*?)(?=Link naar PDF|Graphic|Load-Date|$)', art, re.DOTALL | re.IGNORECASE)
        body = clean_text(body_match.group(1)) if body_match else ""

        if body:
            extracted.append({
                "title": title,
                "date": date,
                "author": author,
                "full_text": body
            })
    return extracted

def run_batch_processor(folder_name, output_name):
    all_data = []
    
    if not os.path.exists(folder_name):
        print(f"Error: Folder '{folder_name}' not found.")
        return

    for filename in os.listdir(folder_name):
        if filename.lower().endswith(".docx") and not filename.startswith("~$"):
            print(f"Processing: {filename}...")
            try:
                raw_text = get_docx_text(os.path.join(folder_name, filename))
                parsed = parse_nrc_batch(raw_text)
                all_data.extend(parsed)
                print(f"   Successfully found {len(parsed)} articles.")
            except Exception as e:
                print(f"   Failed to read {filename}: {e}")

    with open(output_name, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, indent=4, ensure_ascii=False)
    
    print(f"\nFINISH: Total articles extracted: {len(all_data)}")
    print(f"File created: {output_name}")


if __name__ == "__main__":
    run_batch_processor('nrc_files', 'nrc_master_data.json')