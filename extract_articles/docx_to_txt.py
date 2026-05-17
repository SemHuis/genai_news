from docx import Document
import argparse
import os
import sys
import re

def normalize_text(text):
    """Normalize unusual line separators and multiple newlines.
    """
    if not text:
        return ""
    
    # Replace Unicode line/paragraph separators with standard newlines
    text = re.sub(r'[\u2028\u2029\u0085]', '\n', text)
    # Replace carriage returns + newline with just newline
    text = re.sub(r'\r\n', '\n', text)
    # Replace lone carriage returns with newlines
    text = re.sub(r'\r', '\n', text)
    # Replace 3+ consecutive newlines with exactly 2 (preserve paragraph breaks)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip trailing/leading whitespace from each line but preserve structure
    lines = text.split('\n')
    lines = [line.rstrip() for line in lines]
    text = '\n'.join(lines)
    
    return text.strip()

def get_docx_text(path):
    """Extracts text from paragraphs AND tables to ensure nothing is missed."""
    doc = Document(path)
    content = []
    # Extract from paragraphs
    for para in doc.paragraphs:
        normalized = normalize_text(para.text)
        if normalized:  # only add non-empty paragraphs
            content.append(normalized)
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                normalized = normalize_text(cell.text)
                if normalized:
                    content.append(normalized)
    # Join paragraphs with double newlines for readability
    return "\n\n".join(content)

def main():
    parser = argparse.ArgumentParser(description="Extract and clean news articles from DOCX files.")
    parser.add_argument('-d', '--directory', help='Path to the directory containing .DOCX files')
    parser.add_argument('-o', '--output', default='all_articles.txt', help='Output text file (default: all_articles.txt)')
    args = parser.parse_args()

    all_articles = []
    
    # Ensure the directory exists before proceeding
    if not os.path.isdir(args.directory):
        print(f"Error: The directory '{args.directory}' does not exist.")
        sys.exit(1)

    # Process Files
    for filename in os.listdir(args.directory):
        if filename.lower().endswith(".docx"):
            file_path = os.path.join(args.directory, filename)
            try:
                text = get_docx_text(file_path)
                all_articles.append(text)
            except Exception as e:
                print(f"Could not read {filename}: {e}")
    
    with open(args.output, "w", encoding="utf-8") as f:
        for article in all_articles:
            # Separate articles by two newlines
            f.write(article + "\n\n")
            
    print(f"Processed {len(all_articles)} files into {args.output}")

if __name__ == "__main__":    
    main()