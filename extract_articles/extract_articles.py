import os
import re
import json
from docx import Document

def get_docx_text(path):
    """Extracts text from paragraphs AND tables to ensure nothing is missed."""
    doc = Document(path)
    content = []
    # Extract from paragraphs
    for para in doc.paragraphs:
        content.append(para.text)
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Add newline to keep structure
                content.append(cell.text + "\n")
    return "\n".join(content)

def clean_text(text):
    """Removes tags and standardizes whitespace."""
    if not text: return ""
    text = re.sub(r'\\', '', text)
    # Convert multiple spaces/newlines into clean structure
    return text.strip()


def parse_date(raw_text):
    """Parse a date string (often Dutch) and return it as DD/MM/YYYY.

    Examples handled:
      - '26 augustus 2024 maandag' -> '26/08/2024'
      - 'maandag 26 augustus 2024' -> '26/08/2024'
      - '26/08/2024' or '26-08-2024' -> normalized to '26/08/2024'
    If parsing fails, returns the cleaned input unchanged.
    """
    if not raw_text:
        return "Unknown Date"
    s = clean_text(raw_text)

    # Remove common weekday names (Dutch)
    s = re.sub(r"\b(maandag|dinsdag|woensdag|donderdag|vrijdag|zaterdag|zondag)\b", "", s, flags=re.I)

    # Map Dutch month names to month numbers
    months = {
        'januari':1, 'februari':2, 'maart':3, 'april':4, 'mei':5, 'juni':6,
        'juli':7, 'augustus':8, 'september':9, 'oktober':10, 'november':11, 'december':12
    }

    # Try patterns like '26 augustus 2024'
    m = re.search(r"(?P<day>\d{1,2})\s+(?P<month>[A-Za-z]+)\s+(?P<year>\d{4})", s, flags=re.I)
    if m:
        day = int(m.group('day'))
        month_word = m.group('month').lower()
        year = m.group('year')
        month = months.get(month_word)
        if month:
            return f"{day:02d}/{month:02d}/{year}"

    # Try numeric formats like 26/08/2024 or 26-08-2024 or 26.08.2024
    m2 = re.search(r"(?P<day>\d{1,2})[\/\-\.](?P<month>\d{1,2})[\/\-\.](?P<year>\d{2,4})", s)
    if m2:
        day = int(m2.group('day'))
        month = int(m2.group('month'))
        year = m2.group('year')
        if len(year) == 2:
            # assume 2000s for two-digit years (adjust if needed)
            year = '20' + year
        return f"{day:02d}/{month:02d}/{year}"

    # Try pattern '26 augustus' with implicit year (rare) -> leave as-is
    return s


def extract_section(art):
    """Try to extract a section/category for the article.

    Strategy:
      1. Look for explicit labels like 'SECTION:', 'SECTIE:' or 'Rubriek:'
      2. Fallback: check the first few lines for a short all-caps line
         that is likely a section header (e.g. 'BINNENLAND', 'SPORT').
    Returns cleaned section string or 'Unknown'.
    """
    if not art:
        return "Unknown"

    # 1) explicit labels
    for pat in [r'(?im)^\s*SECTION[:\-\s]+(.+)$', r'(?im)^\s*SECTIE[:\-\s]+(.+)$', r'(?im)^\s*Rubriek[:\-\s]+(.+)$']:
        m = re.search(pat, art)
        if m:
            return clean_text(m.group(1))

    # 2) fallback: look at first lines for an uppercase short line
    lines = [l.strip() for l in art.splitlines() if l.strip()]
    for ln in lines[:8]:
        # skip lines that are the title or obvious markers
        if re.search(r'\bNRC\b', ln, flags=re.I):
            continue
        if re.search(r'VOLLEDIGE TEKST', ln, flags=re.I):
            continue
        # consider this a section if it's mostly uppercase letters and short
        words = ln.split()
        if 1 <= len(words) <= 5 and re.fullmatch(r"[A-Z0-9\-\'\.]{2,40}", ln):
            return clean_text(ln.title())

    return "Unknown"

def parse_nrc_batch(raw_text):
    """Parses raw text using multiple possible separators."""
    # Split by common separators (case-insensitive)
    # Handles "End of Document", "END OF DOCUMENT", or "Document 1 of 250"
    articles = re.split(r'(?i)End of Document|Document \d+ of \d+', raw_text)
    
    extracted = []
    for art in articles:
        # require at least the NRC marker (case-insensitive)
        if not re.search(r"\bNRC\b", art, flags=re.I):
            continue

        # 1. Title: Look for text followed by NRC
        # We look for the title at the start of the block or after empty lines
        title_match = re.search(r'(?:^|\n)\s*(\d+\.\s*)?([^\n]+)\n\s*NRC', art, re.DOTALL | re.IGNORECASE)
        title = clean_text(title_match.group(2)) if title_match else "Unknown Title"

        # 2. Date: The line immediately following NRC (case-insensitive)
        date_match = re.search(r'NRC\s*\n\s*([^\n]+)', art, flags=re.I)
        if date_match:
            date = parse_date(date_match.group(1))
        else:
            # fallback: look for a date-like pattern anywhere in the article
            date_guess = re.search(r"\d{1,2}\s+[A-Za-z]+\s+\d{4}", art, flags=re.I)
            if date_guess:
                date = parse_date(date_guess.group(0))
            else:
                date = "Unknown Date"

        # 3. Author: Byline or footer (case-insensitive)
        author_match = re.search(r'Byline:\s*([^\n]+)', art, flags=re.I)
        author = clean_text(author_match.group(1)) if author_match else "Redactie"

        # 4. Body: Prefer VOLLEDIGE TEKST marker, but fall back to everything
        # after the NRC/date line if the marker is missing
        body_match = re.search(r'VOLLEDIGE TEKST:(.*?)(?=Link naar PDF|Graphic|Load-Date|$)', art, re.DOTALL | re.IGNORECASE)
        if body_match:
            body = clean_text(body_match.group(1))
        else:
            # find position after the NRC + optional date line
            m_nrc_line = re.search(r'NRC\s*\n\s*[^\n]+', art, flags=re.I)
            start_pos = m_nrc_line.end() if m_nrc_line else 0
            remainder = art[start_pos:]
            # strip common footers
            body = clean_text(re.split(r'(?i)Link naar PDF|Graphic|Load-Date', remainder)[0])

        if body:
            extracted.append({
                "title": title,
                "date": date,
                "author": author,
                "section": extract_section(art),
                "full_text": body
            })
    return extracted

def run_batch_processor(folder_name, output_name):
    all_data = []
    
    if not os.path.exists(folder_name):
        print(f"Error: Folder '{folder_name}' not found.")
        return

    for filename in os.listdir(folder_name):
        if filename.lower().endswith(".docx") and not filename.startswith("~$"):
            print(f"Processing: {filename}...")
            try:
                raw_text = get_docx_text(os.path.join(folder_name, filename))
                parsed = parse_nrc_batch(raw_text)
                all_data.extend(parsed)
                print(f"   Successfully found {len(parsed)} articles.")
            except Exception as e:
                print(f"   Failed to read {filename}: {e}")

    with open(output_name, 'w', encoding='utf-8') as f:
        json.dump(all_data, f, indent=4, ensure_ascii=False)
    
    print(f"\nFINISH: Total articles extracted: {len(all_data)}")
    print(f"File created: {output_name}")


if __name__ == "__main__":
    run_batch_processor('nrc_files', 'nrc_master_data.json')